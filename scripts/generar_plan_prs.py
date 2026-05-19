from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
OUT_DIR.mkdir(exist_ok=True)
DOCX_PATH = OUT_DIR / "plan_pull_requests_gitflow_manejo_buses.docx"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(8.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_bullet(document, text):
    document.add_paragraph(text, style="List Bullet")


def add_number(document, text):
    document.add_paragraph(text, style="List Number")


def add_code_block(document, lines):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(8)
    for line in lines:
        run = paragraph.add_run(line + "\n")
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(8.5)


def add_table(document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(header_cells[idx], header, bold=True)
        shade_cell(header_cells[idx], "D9EAF7")
        if widths:
            header_cells[idx].width = widths[idx]

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], str(value))
            if widths:
                cells[idx].width = widths[idx]

    document.add_paragraph()
    return table


def setup_styles(document):
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Normal"].font.size = Pt(10)

    for name, size, color in [
        ("Title", 20, "0F172A"),
        ("Heading 1", 15, "0F172A"),
        ("Heading 2", 12, "1E3A5F"),
        ("Heading 3", 10.5, "334155"),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)


def build_document():
    document = Document()
    setup_styles(document)

    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Plan de Pull Requests con Git Flow")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(15, 23, 42)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Proyecto: Manejo Buses - Gestion y venta de pasajes interprovinciales").italic = True

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Rama base: develop | Equipo: 2 personas | Flujo: Git Flow")

    document.add_heading("1. Objetivo del documento", level=1)
    document.add_paragraph(
        "Este documento divide el desarrollo completo del proyecto en Pull Requests manejables para dos personas. "
        "La division esta pensada para trabajar desde la rama develop usando Git Flow, evidenciar colaboracion, "
        "mantener cambios revisables y facilitar la auditoria de configuracion solicitada en el segundo parcial."
    )

    document.add_heading("2. Reglas generales de Git Flow", level=1)
    add_bullet(document, "La rama develop es la base de integracion diaria.")
    add_bullet(document, "Cada tarea se trabaja en una rama feature/* creada desde develop.")
    add_bullet(document, "Cada PR debe apuntar hacia develop, nunca directo a main.")
    add_bullet(document, "Antes de abrir PR se debe ejecutar php artisan test y, si cambia frontend, npm run build.")
    add_bullet(document, "Cada PR debe estar vinculado a una tarea o cambio en Jira.")
    add_bullet(document, "La otra persona revisa el PR antes de hacer merge.")
    add_bullet(document, "main se usa solo para versiones estables, mediante release/* y hotfix/*.")

    document.add_heading("3. Comandos base para cualquier PR", level=1)
    add_code_block(
        document,
        [
            "git checkout develop",
            "git pull origin develop",
            "git checkout -b feature/nombre-de-la-tarea",
            "# realizar cambios",
            "php artisan test",
            "npm run build",
            "git status",
            "git add .",
            "git commit -m \"feat: descripcion corta de la tarea\"",
            "git push -u origin feature/nombre-de-la-tarea",
            "# abrir Pull Request hacia develop en GitHub",
        ],
    )

    document.add_heading("4. Convencion de ramas y commits", level=1)
    add_table(
        document,
        ["Tipo", "Formato", "Ejemplo"],
        [
            ["Feature", "feature/modulo-descripcion", "feature/crud-buses"],
            ["Fix", "fix/modulo-problema", "fix/login-validacion-activo"],
            ["Release", "release/vX.Y.Z", "release/v1.0.0"],
            ["Hotfix", "hotfix/problema-produccion", "hotfix/error-migracion-rutas"],
            ["Commit", "tipo: mensaje breve", "feat: implementar CRUD de buses"],
        ],
        [Inches(1.1), Inches(2.2), Inches(3.3)],
    )

    document.add_heading("5. Division general por persona", level=1)
    add_table(
        document,
        ["Persona", "Responsabilidad principal", "Responsabilidad secundaria"],
        [
            [
                "Persona A",
                "Base tecnica, autenticacion, usuarios internos, cooperativas, buses, rutas/frecuencias, panel administrativo.",
                "Revisar PRs de ventas, pagos, boletos y notificaciones.",
            ],
            [
                "Persona B",
                "Interfaz publica, busqueda de viajes, compra de boletos, pagos, comprobantes, historial y validacion de acceso.",
                "Revisar PRs de configuracion, modelos, migraciones y administracion.",
            ],
        ],
        [Inches(1.0), Inches(3.2), Inches(3.2)],
    )

    document.add_heading("6. Roadmap detallado por Pull Request", level=1)
    prs = [
        ["PR-01", "A", "feature/base-configuracion-db", "Base Laravel, .env.example, conexion MySQL, migraciones y modelos iniciales.", "Proyecto corre, migraciones ejecutan, php artisan test pasa."],
        ["PR-02", "B", "feature/interfaz-principal", "Layout principal, menu publico, pagina de inicio de buses y estructura visual general.", "Inicio separado del login, menu responsive, pagina publica en /."],
        ["PR-03", "A", "feature/autenticacion-usuarios", "Registro, login, logout, dashboard protegido y roles base.", "Usuarios nuevos quedan como cliente, rutas protegidas por auth."],
        ["PR-04", "A", "feature/crud-cooperativas", "CRUD de cooperativas con datos de contacto, logo path y estado activo.", "Se puede crear, listar, editar, ver y eliminar cooperativas."],
        ["PR-05", "B", "feature/crud-ciudades", "CRUD de ciudades y provincias para origen, destino y paradas.", "Ciudades activas disponibles para rutas y frecuencias."],
        ["PR-06", "A", "feature/crud-buses", "CRUD de buses asociado a cooperativas.", "Numero unico por cooperativa, placa unica, filtros basicos."],
        ["PR-07", "B", "feature/tipos-asientos-asientos", "CRUD de tipos de asiento y asientos por bus.", "Buses tienen configuracion de asientos Normal/VIP y capacidad consistente."],
        ["PR-08", "A", "feature/crud-rutas", "CRUD de rutas relacionado con cooperativas, buses, origen y destino.", "Ruta valida origen distinto de destino y bus opcional."],
        ["PR-09", "A", "feature/frecuencias-ant-paradas", "CRUD de frecuencias ANT y paradas intermedias ordenadas.", "Resolucion ANT, hora, estado y paradas en orden."],
        ["PR-10", "B", "feature/salidas-hoja-ruta", "Generacion manual de salidas/hoja de ruta desde frecuencias habilitadas.", "Salida asigna bus, fecha, hora, precio base y estado."],
        ["PR-11", "B", "feature/busqueda-viajes", "Pantalla publica para buscar destinos y frecuencias disponibles.", "Filtros por cooperativa, tipo asiento, chasis/carroceria y tipo de viaje."],
        ["PR-12", "B", "feature/seleccion-asientos", "Seleccion de asientos disponibles por salida y tipo.", "No permite vender asiento ocupado en la misma salida."],
        ["PR-13", "B", "feature/venta-boletos-cliente", "Compra de boletos para cliente final con descuentos.", "Registra boleto con origen/destino, pasajero, asiento, precio y estado."],
        ["PR-14", "A", "feature/venta-boletos-oficinista", "Venta interna para rol oficinista.", "Oficinista puede vender/validar boletos desde panel protegido."],
        ["PR-15", "B", "feature/pagos-comprobantes", "Metodo deposito/transferencia y carga de comprobante.", "Pago queda pendiente y puede validarse manualmente."],
        ["PR-16", "A", "feature/validacion-pagos", "Validacion/rechazo de pagos por oficinista.", "Actualiza estado del pago y del boleto segun validacion."],
        ["PR-17", "B", "feature/boleto-pdf-codigo", "Generar boleto descargable con codigo QR o codigo de barras.", "Boleto visible despues de compra y desde historial."],
        ["PR-18", "B", "feature/historial-compras", "Historial de compras para usuario cliente.", "Cliente ve boletos, estados, pagos y descarga boleto."],
        ["PR-19", "A", "feature/acceso-pasajeros", "Modulo para chofer/ayudante: escaneo o ingreso de codigo.", "Registra acceso permitido/rechazado en registro_accesos."],
        ["PR-20", "A", "feature/configuracion-aplicacion", "CRUD/configuracion de logo, colores, redes, soporte.", "Configuracion se refleja en layout cuando aplique."],
        ["PR-21", "A", "feature/roles-permisos", "Middleware y restricciones por rol.", "Cooperativa gestiona buses propios, oficinista ventas, cliente compras."],
        ["PR-22", "B", "feature/notificaciones-correo", "Notificaciones por correo o app en pasos relevantes.", "Compra, pago pendiente, pago validado y boleto emitido notifican."],
        ["PR-23", "A+B", "feature/reportes-auditoria", "Reportes para evidenciar trabajo, cambios y estados del sistema.", "Dashboard o vistas con resumen de ventas, pagos, accesos y auditoria."],
        ["PR-24", "A+B", "release/v1.0.0", "Integracion final, QA, documentacion, informe y preparacion de presentacion.", "Release estable en main con tag y pruebas finales."],
    ]
    add_table(document, ["PR", "Resp.", "Rama", "Alcance", "Criterio de aceptacion"], prs)

    document.add_heading("7. Orden recomendado de trabajo", level=1)
    order = [
        "Semana 1: PR-01, PR-02, PR-03.",
        "Semana 2: PR-04, PR-05, PR-06, PR-07.",
        "Semana 3: PR-08, PR-09, PR-10.",
        "Semana 4: PR-11, PR-12, PR-13, PR-14.",
        "Semana 5: PR-15, PR-16, PR-17, PR-18.",
        "Semana 6: PR-19, PR-20, PR-21, PR-22.",
        "Semana 7: PR-23, PR-24, informe, evidencias, presentacion y demo.",
    ]
    for item in order:
        add_number(document, item)

    document.add_heading("8. Detalle por PR y comandos sugeridos", level=1)
    for pr, person, branch, scope, acceptance in prs:
        document.add_heading(f"{pr} - {branch}", level=2)
        add_bullet(document, f"Responsable: {person}.")
        add_bullet(document, f"Alcance: {scope}")
        add_bullet(document, f"Criterio de aceptacion: {acceptance}")
        add_code_block(
            document,
            [
                "git checkout develop",
                "git pull origin develop",
                f"git checkout -b {branch}",
                "# implementar cambios del alcance",
                "php artisan test",
                "npm run build",
                "git add .",
                f"git commit -m \"feat: {scope[:55].lower()}\"",
                f"git push -u origin {branch}",
            ],
        )

    document.add_heading("9. Plantilla sugerida para cada Pull Request", level=1)
    add_code_block(
        document,
        [
            "## Descripcion",
            "Resumen breve del cambio implementado.",
            "",
            "## Cambios realizados",
            "- Cambio 1",
            "- Cambio 2",
            "",
            "## Validaciones",
            "- php artisan test",
            "- npm run build",
            "",
            "## Evidencias",
            "- Capturas de pantalla",
            "- Ruta probada",
            "- Issue/Jira asociado",
            "",
            "## Checklist",
            "- [ ] Codigo revisado localmente",
            "- [ ] Migraciones probadas si aplica",
            "- [ ] No se mezclan tareas no relacionadas",
            "- [ ] PR apunta a develop",
        ],
    )

    document.add_heading("10. Flujo de control de cambios con Jira", level=1)
    add_table(
        document,
        ["Estado", "Responsable", "Descripcion"],
        [
            ["Backlog", "Comite / equipo", "Solicitud de cambio registrada y priorizada."],
            ["To Do", "Responsable asignado", "Cambio aprobado para implementacion."],
            ["In Progress", "Desarrollador", "Rama feature creada desde develop y trabajo en curso."],
            ["In Review", "Revisor", "Pull Request abierto, pruebas y evidencias adjuntas."],
            ["Approved", "Comite / revisor", "Cambio aceptado para merge a develop."],
            ["Done", "Equipo", "PR mergeado, rama eliminada y Jira actualizado."],
        ],
        [Inches(1.2), Inches(1.7), Inches(4.5)],
    )

    document.add_heading("11. Comandos para release final", level=1)
    add_code_block(
        document,
        [
            "git checkout develop",
            "git pull origin develop",
            "git checkout -b release/v1.0.0",
            "php artisan test",
            "npm run build",
            "git add .",
            "git commit -m \"chore: preparar release v1.0.0\"",
            "git push -u origin release/v1.0.0",
            "# abrir PR release/v1.0.0 -> main",
            "# luego abrir PR main -> develop para sincronizar",
            "git tag -a v1.0.0 -m \"Version estable segundo parcial\"",
            "git push origin v1.0.0",
        ],
    )

    document.add_heading("12. Checklist final del proyecto", level=1)
    final_checks = [
        "Repositorio en GitHub con ramas, commits y Pull Requests.",
        "Issues o tareas en Jira enlazadas a cada PR.",
        "Evidencia de revisiones entre integrantes.",
        "Migraciones y seeders necesarios para demo.",
        "Autenticacion y roles base funcionando.",
        "CRUDs principales funcionando: cooperativas, ciudades, buses, rutas, frecuencias y configuracion.",
        "Flujo de compra de boletos y pagos funcionando.",
        "Validacion de acceso de pasajeros funcionando.",
        "Informe con capturas del flujo Git, Jira y demo de la aplicacion.",
        "Presentacion grupal con demo de 20 a 30 minutos.",
    ]
    for item in final_checks:
        add_bullet(document, item)

    document.add_heading("13. Nota de alcance", level=1)
    document.add_paragraph(
        "La division propone PRs pequenos y revisables. Si alguna funcionalidad resulta demasiado grande, se recomienda dividirla "
        "en sub-PRs manteniendo el mismo prefijo de modulo. Por ejemplo, feature/venta-boletos-cliente-formulario y "
        "feature/venta-boletos-cliente-confirmacion."
    )

    document.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build_document())
